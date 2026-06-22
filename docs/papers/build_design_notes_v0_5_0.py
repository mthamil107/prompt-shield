"""Build the v0.5.0 design-notes DOCX for arXiv / Zenodo upload.

Produces ``docs/papers/design-notes-v0.5.0.docx``. Source-of-truth prose lives
in ``docs/design-notes-v0.5.0.md`` (CC BY 4.0) — this script is a one-way
formatter that wraps that text in academic-paper styling so arXiv / Zenodo
get a clean PDF on upload.

Reproduce any time by running this script; do NOT hand-edit the generated
DOCX. To convert to PDF on Windows: open in Microsoft Word, ``File → Save As
→ PDF``. Then upload the PDF to arXiv (cs.CR) and Zenodo.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).parent.parent.parent
SRC_MD = REPO_ROOT / "docs" / "design-notes-v0.5.0.md"
OUT_DOCX = REPO_ROOT / "docs" / "papers" / "design-notes-v0.5.0.docx"

TITLE = (
    "Design Notes: Seven Cross-Domain Pre-processing and Detection Techniques "
    "for Prompt-Injection Defense (prompt-shield v0.5.0)"
)
AUTHOR = "Thamilvendhan Munirathinam"
AFFILIATION = "Independent — https://github.com/mthamil107/prompt-shield"
DATE_STR = "2026-06-18"
ABSTRACT = (
    "This technical report documents the novel detection and pre-processing "
    "techniques introduced in version 0.5.0 of the prompt-shield library — "
    "an open-source prompt-injection firewall released under Apache 2.0. "
    "Seven techniques are described with sufficient algorithmic detail to "
    "be independently reimplemented: many-shot structural detection via "
    "coefficient-of-variation analysis of turn markers (d029); a fail-soft "
    "YAML-rules engine with highest-severity-wins precedence (d030); a "
    "two-stage language-enforcement detector combining Unicode script-ratio "
    "analysis with optional langdetect (d031); operator-defined denied-topic "
    "detection with multi-group keyword clusters and most-hits-wins "
    "resolution (d032); multi-turn topic-drift detection via Jaccard "
    "similarity between the current turn and a self-anchored fingerprint "
    "of the conversation's first turns (d033); a four-stage idempotent "
    "normalization pipeline that emits change-tracking metadata; and a "
    "fan-out multi-encoding preprocessor (base64, hex, URL, HTML entities, "
    "ROT13) that feeds decoded candidates back through the full detector "
    "stack. We also document the removal of the 512-token input-length cap "
    "on the DeBERTa-v3 semantic classifier via overlap-chunked max-pool "
    "aggregation. This is a companion to arXiv:2604.18248 (the main "
    "prompt-shield paper covering d027 / d028 / adversarial fatigue). It is "
    "published as a dated public disclosure under CC BY 4.0; the techniques "
    "described here are released as prior art and the author makes no claim "
    "to patent rights over them."
)
KEYWORDS = (
    "prompt injection, LLM security, AI safety, prior art, "
    "many-shot jailbreak, topic drift, homoglyph normalization, "
    "multi-encoding decoder, language enforcement"
)


def _add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(15)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run(AUTHOR).bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run(AFFILIATION).italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.add_run(DATE_STR)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("Abstract. ").bold = True
    p.add_run(ABSTRACT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.add_run("Keywords: ").bold = True
    p.add_run(KEYWORDS)


def _add_para(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if bold_lead:
        p.add_run(bold_lead).bold = True
        p.add_run(" ")
    p.add_run(text)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run(text)
    r.bold = True
    r.font.size = Pt(13 if level == 1 else 11)


def _code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)


def build() -> None:
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    _add_title_block(doc)

    # --- Section 1 ---
    _heading(doc, "1. Purpose and scope", 1)
    _add_para(
        doc,
        "This document is a technical design note for the novel detection "
        "and pre-processing techniques introduced in prompt-shield v0.5.0 "
        "that are not covered in the companion arXiv paper (arXiv:2604.18248). "
        "It is published as a dated public disclosure so that the techniques "
        "described here are unambiguously prior art as of the timestamp on "
        "this document and its DOI on Zenodo.",
    )
    _add_para(
        doc,
        "The companion paper covers d027 (stylometric discontinuity), d028 "
        "(Smith-Waterman sequence alignment with semantic substitution "
        "matrix), and the adversarial fatigue tracker. This document covers "
        "the seven additional cross-domain techniques shipped in v0.5.0, in "
        "enough algorithmic detail to be independently reimplemented from "
        "this text alone.",
    )
    _add_para(
        doc,
        "All techniques are released under Apache 2.0 (code) and CC BY 4.0 "
        "(this description). The intent of this disclosure is not to "
        "restrict use — it is to ensure that no future patent application "
        "can claim novelty over these techniques.",
    )

    # --- Section 2: d029 ---
    _heading(doc, "2. d029 — Many-shot Structural Detection", 1)
    _add_para(
        doc,
        "Anthropic (2024) disclosed a class of attack where a prompt contains "
        "a large number of structurally identical fake conversation turns "
        "that condition the model to comply with the final turn. Verbatim "
        "regex on any single turn fails. Structural features detect it:",
        bold_lead="Threat model.",
    )
    _bullet(
        doc,
        "Turn marker density: count canonical markers (Q:, A:, Human:, "
        "Assistant:, [INST], etc.) divided by total non-whitespace chars. "
        "Above ~1 marker per 80 chars in long prompts is anomalous.",
    )
    _bullet(
        doc,
        "Inter-turn period regularity: measure character distance between "
        "adjacent markers and compute coefficient of variation (stdev/mean). "
        "Genuine conversations have CV ≥ 0.5; mechanically generated "
        "many-shot attacks have CV ≤ 0.15.",
    )
    _bullet(
        doc,
        "Total turn count threshold: ≥ 30 turn markers in a single prompt "
        "is rare in legitimate use.",
    )
    _add_para(
        doc,
        "Detection fires when turn count ≥ 30 AND turn density above "
        "threshold AND CV ≤ 0.15. Confidence scales with how far regularity "
        "exceeds threshold.",
        bold_lead="Decision rule.",
    )
    _add_para(
        doc,
        "Existing literature on many-shot jailbreaking focuses on attack "
        "construction. This technique reframes detection as periodicity "
        "analysis on structural markers, borrowing from time-series anomaly "
        "detection — short conversations have irregular intervals, "
        "mechanically padded ones do not.",
        bold_lead="Novelty claim.",
    )
    _add_para(
        doc,
        "src/prompt_shield/detectors/d029_many_shot_structural.py.",
        bold_lead="Implementation reference.",
    )

    # --- Section 3: d030 ---
    _heading(doc, "3. d030 — Custom YAML Rules Engine", 1)
    _add_para(
        doc,
        "Operators routinely need to block organization-specific terms "
        "(internal codenames, sensitive data class names) that no general "
        "detector can know. Hand-editing detector Python is high-friction; "
        "a declarative rule format lowers it.",
        bold_lead="Threat model.",
    )
    _add_para(
        doc,
        "A directory of YAML rule files is loaded at engine startup. Each "
        "rule has the schema:",
        bold_lead="Technique.",
    )
    _code_block(
        doc,
        "rules:\n"
        "  - id: internal-codename\n"
        "    pattern: '\\binternal-codename-X\\b'\n"
        "    severity: high          # critical | high | medium | low\n"
        "    action: block           # block | flag | log\n"
        "    description: \"...\"\n"
        "    case_sensitive: false   # optional, default false",
    )
    _add_para(
        doc,
        "Loading is fail-soft per rule: malformed YAML, invalid regex, or "
        "entries missing id/pattern are logged and skipped without aborting "
        "the rest of the ruleset. On detection, when multiple rules match, "
        "highest-severity wins (critical > high > medium > low). The chosen "
        "rule's severity, action, and metadata propagate to the engine-level "
        "decision.",
    )
    _add_para(
        doc,
        "The fail-soft rule loader + highest-severity-wins precedence + "
        "per-rule action verb is a distinctive combination — most regex-rule "
        "engines are either fail-hard or first-match-wins.",
        bold_lead="Novelty claim.",
    )
    _add_para(
        doc,
        "src/prompt_shield/detectors/d030_custom_rules.py.",
        bold_lead="Implementation reference.",
    )

    # --- Section 4: d031 ---
    _heading(doc, "4. d031 — Two-Stage Language Enforcement", 1)
    _add_para(
        doc,
        "Many deployments are explicitly English-only or restricted to a "
        "small allow-list. Multilingual jailbreaks (translations of common "
        "attack templates into less-trained languages) bypass English-only "
        "RLHF safety signal. Blocking non-allow-list languages is a cheap, "
        "high-precision input gate.",
        bold_lead="Threat model.",
    )
    _add_para(
        doc,
        "Detection is two-stage to minimize dependency surface.",
        bold_lead="Technique.",
    )
    _add_para(
        doc,
        "Pre-compiled regexes for Unicode script blocks: Cyrillic, Greek, "
        "Arabic, Hebrew, Devanagari, Thai, and the CJK union (Han, Hiragana, "
        "Katakana, Hangul). For each script, compute the ratio of matching "
        "characters to total non-whitespace, non-digit characters. If any "
        "non-allowed script exceeds 15% of input characters, the inferred "
        "language is rejected immediately.",
        bold_lead="Stage 1 — fast-path script analysis (no dependencies).",
    )
    _add_para(
        doc,
        "If the input is Latin-script and the fast path doesn't trigger, "
        "optionally invoke langdetect with seed=0 for determinism. Accept "
        "the top guess only if probability ≥ 0.75; otherwise treat as "
        "allowed (high uncertainty defaults to permissive to avoid false "
        "positives on short or technical input).",
        bold_lead="Stage 2 — langdetect fallback (optional dependency).",
    )
    _add_para(
        doc,
        "The fast-path script-ratio gate is well-known in isolation. The "
        "novelty is the two-stage gated combination: script-fraction analysis "
        "as a no-dependency fast path, with langdetect invoked only when "
        "the fast path is inconclusive AND optional dependencies are present. "
        "Makes the detector deployable in environments that cannot pull in "
        "langdetect while still benefiting from it when available.",
        bold_lead="Novelty claim.",
    )
    _add_para(
        doc,
        "src/prompt_shield/detectors/d031_language_enforcement.py.",
        bold_lead="Implementation reference.",
    )

    # --- Section 5: d032 ---
    _heading(doc, "5. d032 — Operator-Defined Denied Topics", 1)
    _add_para(
        doc,
        "Deployments often need to block off-topic requests entirely "
        "(medical, legal, political opinions in a code-assistant). These are "
        "not 'attacks' in the prompt-injection sense — they are policy "
        "violations. The same engine should evaluate them with a single "
        "declarative interface.",
        bold_lead="Threat model.",
    )
    _add_para(
        doc,
        "Operator-defined topic groups, each a named keyword cluster. A "
        "detection fires when the number of keyword hits from any single "
        "group meets a configurable minimum threshold (default 2).",
        bold_lead="Technique.",
    )
    _code_block(
        doc,
        "d032_topic_enforcement:\n"
        "  denied_topics:\n"
        "    - name: medical_advice\n"
        "      keywords: [\"diagnose\", \"prescription\", \"dosage\", \"symptoms\"]\n"
        "      severity: high\n"
        "    - name: legal_advice\n"
        "      keywords: [\"lawsuit\", \"attorney\", \"court\", \"litigation\"]\n"
        "      severity: medium\n"
        "  min_keyword_hits: 2",
    )
    _add_para(
        doc,
        "Keywords are matched as word-bounded literal phrases. Among multiple "
        "matching topics, the topic with the most hits wins; ties resolved by "
        "configuration order.",
    )
    _add_para(
        doc,
        "Combination of (a) multiple competing keyword groups with per-group "
        "severity, (b) min-hits-per-group threshold (so isolated word usage "
        "doesn't fire — \"I'm not asking for medical advice\" doesn't trigger "
        "on \"advice\"), and (c) most-hits-wins precedence. This sits between "
        "naive single-keyword block-lists and full ML topic classifiers.",
        bold_lead="Novelty claim.",
    )
    _add_para(
        doc,
        "src/prompt_shield/detectors/d032_topic_enforcement.py.",
        bold_lead="Implementation reference.",
    )

    # --- Section 6: d033 ---
    _heading(
        doc,
        "6. d033 — Multi-Turn Topic Drift via Jaccard Anchor Similarity",
        1,
    )
    _add_para(
        doc,
        "A class of slow jailbreaks operates across many turns. Each "
        "individual turn is benign — a coding question, then a casual aside, "
        "then a hypothetical, then an escalated request. No single turn is "
        "detected; the cumulative drift from the conversation's established "
        "topic is what makes the final harmful request succeed.",
        bold_lead="Threat model.",
    )
    _add_para(
        doc,
        "Anchored n-gram drift.",
        bold_lead="Technique.",
    )
    _bullet(
        doc,
        "Build the anchor: concatenate the first N turns of the conversation "
        "(default N=2). This is the 'what the conversation is about' "
        "reference. Computed once per session.",
    )
    _bullet(
        doc,
        "Tokenize and stopword-filter both the anchor and the current turn. "
        "Use a small English stoplist (~70 common words) to remove function "
        "words that would inflate similarity.",
    )
    _bullet(
        doc,
        "N-gram fingerprint: build the set of token n-grams (default bigrams, "
        "configurable). Each fingerprint is a set of tuples.",
    )
    _bullet(
        doc,
        "Jaccard similarity between the anchor fingerprint A and current-turn "
        "fingerprint C: |A ∩ C| / |A ∪ C|.",
    )
    _bullet(
        doc,
        "Decision: if similarity is below min_anchor_similarity (default 0.05) "
        "AND the conversation has at least min_turns (default 4), fire "
        "detection. The minimum-turns gate prevents firing on short "
        "interactions that haven't established a topic.",
    )
    _add_para(
        doc,
        "min(0.95, 0.4 + (threshold − similarity) × 4). A complete topic "
        "flip (similarity = 0) yields high confidence; a borderline drift "
        "yields lower confidence.",
        bold_lead="Confidence scaling.",
    )
    _add_para(
        doc,
        "This is the most distinctive technique in v0.5.0. Existing "
        "multi-turn jailbreak detection focuses on either (a) "
        "semantic-embedding distance from a seed prompt (heavy ML "
        "dependency) or (b) detecting specific escalation patterns. The "
        "Jaccard-against-anchor approach is novel in three ways: "
        "(1) no model dependency — pure n-gram set arithmetic, "
        "sub-millisecond evaluation; (2) the anchor is the conversation's "
        "own first turns, not an externally supplied 'safe' prompt — the "
        "detector adapts to each conversation's established topic without "
        "operator configuration of what the topic is; (3) two-gate decision "
        "(similarity < T AND turn_count ≥ M) prevents false-positives on "
        "legitimate topic switches in short conversations.",
        bold_lead="Novelty claim.",
    )
    _add_para(
        doc,
        "src/prompt_shield/detectors/d033_topic_drift.py. Test coverage: "
        "tests/detectors/test_d033_topic_drift.py (10 tests).",
        bold_lead="Implementation reference.",
    )

    # --- Section 7: Normalization ---
    _heading(doc, "7. Normalization Pipeline (pre-detector)", 1)
    _add_para(
        doc,
        "Many obfuscation attacks rely on the input visually rendering "
        "identically to a known attack while being lexically different — "
        "Cyrillic homoglyphs, zero-width characters splitting trigger words, "
        "full-width Unicode reading identically to ASCII. Detectors that "
        "match against the raw byte stream miss these.",
        bold_lead="Threat model.",
    )
    _add_para(
        doc,
        "Four idempotent stages, each independently togglable: "
        "(1) NFKC normalization composes combining marks and maps "
        "compatibility characters; (2) zero-width stripping removes U+200B, "
        "U+200C, U+200D, U+FEFF, U+2060; (3) Cyrillic-to-Latin homoglyph "
        "mapping uses a curated table of 30+ visually-identical "
        "Cyrillic→Latin mappings (one-way only — never reverse, to avoid "
        "mangling legitimate Cyrillic text); (4) whitespace collapse trims "
        "multi-space, tabs, newlines to single spaces.",
        bold_lead="Pipeline structure.",
    )
    _add_para(
        doc,
        "Each stage is idempotent: running it twice produces the same output "
        "as running it once. This guarantees consistent input across "
        "normalization-aware and normalization-naive detectors composed "
        "in the same engine.",
    )
    _add_para(
        doc,
        "The pipeline returns NormalizationResult(text, original, "
        "changes: list[str]) so downstream detectors can see which "
        "normalizations fired — useful both for explanation and as a "
        "detection signal in itself (an input that was zero-width-stripped "
        "is itself suspect).",
        bold_lead="Result object.",
    )
    _add_para(
        doc,
        "The idempotent-stage pipeline with change-tracking output. "
        "Existing libraries typically apply normalizations destructively "
        "without surfacing what changed. The change-tracking output enables "
        "a meta-detector ('input contained zero-width characters') on top "
        "of the standard detectors.",
        bold_lead="Novelty claim.",
    )
    _add_para(
        doc,
        "src/prompt_shield/normalization/pipeline.py.",
        bold_lead="Implementation reference.",
    )

    # --- Section 8: Multi-encoding ---
    _heading(doc, "8. Multi-Encoding Preprocessor (pre-detector)", 1)
    _add_para(
        doc,
        "Encoded payloads are the canonical way to smuggle attack content "
        "past detectors: base64, hex, URL-encoded, HTML-entity-encoded, "
        "ROT13. Single-encoding detectors miss combined attacks "
        "(base64(rot13('ignore...'))).",
        bold_lead="Threat model.",
    )
    _add_para(
        doc,
        "Fan-out candidate set. The preprocessor returns a DecodedSet — a "
        "list of DecodedCandidate(text, encoding, source_span) for each "
        "successfully decoded substring. Detectors then run on each "
        "candidate independently.",
        bold_lead="Technique.",
    )
    _bullet(
        doc,
        "Base64: regex (?<![A-Za-z0-9+/=])([A-Za-z0-9+/]{16,}={0,2})"
        "(?![A-Za-z0-9+/]) — captures contiguous base64-charset runs of "
        "≥ 16 chars with lookaround to preserve padding = characters.",
    )
    _bullet(
        doc,
        "Hex: (?<![0-9a-fA-F])((?:[0-9a-fA-F]{2}){6,})(?![0-9a-fA-F]) — "
        "even-length hex runs of ≥ 12 characters.",
    )
    _bullet(doc, "URL: urllib.parse.unquote applied to %XX-containing substrings.")
    _bullet(doc, "HTML entities: html.unescape applied to &XXX; / &#NNN; patterns.")
    _bullet(
        doc,
        "ROT13: heuristic — any contiguous letters-only word ≥ 7 chars is "
        "a candidate; decoded variant is added.",
    )
    _add_para(
        doc,
        "Decoded candidates feed back through the full detector pipeline. "
        "A detection on a decoded candidate is reported with the original "
        "(encoded) span as the match position.",
    )
    _add_para(
        doc,
        "The fan-out candidate-set model. Most detectors are written as "
        "decode_if_possible(text) → detect(text), handling one encoding per "
        "substring. The fan-out model treats decoded variants as additional "
        "inputs to the full detector stack, so a base64-encoded ROT13-encoded "
        "attack is caught when (a) base64 decode produces ROT13 text, "
        "(b) ROT13 decode produces the original attack, (c) any detector "
        "matches the final decoded form.",
        bold_lead="Novelty claim.",
    )
    _add_para(
        doc,
        "src/prompt_shield/decoders/preprocessor.py.",
        bold_lead="Implementation reference.",
    )

    # --- Section 9: d022 chunking ---
    _heading(doc, "9. Removal of d022 Input-Length Cap (semantic ML classifier)", 1)
    _add_para(
        doc,
        "The DeBERTa-v3 semantic classifier has a 512-token model context. "
        "Inputs longer than ~6000 characters were previously truncated to "
        "the first 512 tokens, leaving the tail of long inputs uncovered.",
        bold_lead="Prior limitation.",
    )
    _add_para(
        doc,
        "Sliding chunked max-pool. Long inputs are split into overlapping "
        "chunks (chunk_size=512 tokens, chunk_stride=384 tokens, "
        "128-token overlap), capped at max_chunks=8. Each chunk is "
        "independently scored by DeBERTa. The final confidence is the max "
        "over all chunk confidences (max-pool aggregation), not the mean.",
        bold_lead="Technique.",
    )
    _add_para(
        doc,
        "Mean-pool dilutes the signal — a 5000-token prompt with one "
        "malicious 100-token segment averages out to 'looks fine'. Max-pool "
        "preserves the signal: any segment scoring high triggers the engine. "
        "For a defensive detector, max-pool is the correct aggregation.",
        bold_lead="Why max-pool.",
    )
    _add_para(
        doc,
        "Chunking itself is standard. The combination of (a) overlap to "
        "avoid splitting tokens across chunks, (b) max-pool aggregation, "
        "and (c) a hard chunk cap (max_chunks=8) to bound worst-case "
        "compute is the contribution. The cap prevents adversarial "
        "input-padding attacks that would force the detector to run for "
        "unbounded time.",
        bold_lead="Novelty claim.",
    )
    _add_para(
        doc,
        "src/prompt_shield/detectors/d022_semantic_classifier.py.",
        bold_lead="Implementation reference.",
    )

    # --- Summary table ---
    _heading(doc, "10. Summary of novelty claims", 1)
    rows = [
        ("Technique", "Novel?", "Primary novelty claim"),
        ("d029 many-shot structural", "Yes", "Periodicity / CV analysis on turn markers"),
        ("d030 custom YAML rules", "Partial", "Fail-soft loader + highest-severity-wins"),
        ("d031 language enforcement", "Yes", "Two-stage script-ratio fast path + optional langdetect"),
        ("d032 topic enforcement", "Yes", "Multi-group keyword clusters + min-hits + most-hits-wins"),
        ("d033 topic drift", "Yes (strongest)", "Jaccard-against-self-anchor + two-gate decision"),
        ("Normalization pipeline", "Partial", "Idempotent stages with change-tracking output"),
        ("Multi-encoding preprocessor", "Yes", "Fan-out candidate set fed back through full detector stack"),
        ("d022 chunking", "Partial", "Overlap + max-pool + hard chunk cap"),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    table.style = "Light List Accent 1"
    for i, (a, b, c) in enumerate(rows):
        cells = table.rows[i].cells
        for j, val in enumerate([a, b, c]):
            cells[j].text = val
            if i == 0:
                for p in cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True

    # --- Disclosure intent ---
    _heading(doc, "11. Disclosure intent", 1)
    _add_para(
        doc,
        "This document, the companion paper arXiv:2604.18248, the public "
        "git history at https://github.com/mthamil107/prompt-shield, and the "
        "released artifact prompt-shield-ai==0.5.0 on PyPI constitute public "
        "disclosure of the techniques described above. They are released "
        "under permissive licenses (CC BY 4.0 for the text, Apache 2.0 for "
        "the code, with the explicit patent grant therein) so that anyone "
        "may freely use, modify, extend, or commercially deploy them.",
    )
    _add_para(
        doc,
        "The author makes no claim of patent rights on these techniques and "
        "disclaims any intent to seek such rights. This disclosure is "
        "intended to establish prior art so that no third party may seek "
        "patent rights over them either.",
    )

    # --- References ---
    _heading(doc, "References", 1)
    _bullet(
        doc,
        "Munirathinam, T. (2026). Beyond Pattern Matching: Seven "
        "Cross-Domain Techniques for Prompt Injection Detection. "
        "arXiv:2604.18248. https://arxiv.org/abs/2604.18248",
    )
    _bullet(
        doc,
        "Anthropic (2024). Many-shot jailbreaking. "
        "https://www.anthropic.com/research/many-shot-jailbreaking",
    )
    _bullet(
        doc,
        "prompt-shield v0.5.0 source code. "
        "https://github.com/mthamil107/prompt-shield/tree/v0.5.0 (Apache 2.0)",
    )
    _bullet(
        doc,
        "CHANGELOG.md, prompt-shield v0.5.0 entry. "
        "https://github.com/mthamil107/prompt-shield/blob/v0.5.0/CHANGELOG.md",
    )

    doc.save(str(OUT_DOCX))
    print(f"Wrote {OUT_DOCX}")


def to_pdf() -> None:
    """Convert the DOCX to PDF using Word (Win/Mac via docx2pdf).

    Requires Microsoft Word installed. On Linux / CI, install LibreOffice
    and run ``soffice --headless --convert-to pdf design-notes-v0.5.0.docx``
    instead.
    """
    try:
        from docx2pdf import convert
    except ImportError:
        print("docx2pdf not installed. Run: pip install docx2pdf")
        return
    out_pdf = OUT_DOCX.with_suffix(".pdf")
    convert(str(OUT_DOCX), str(out_pdf))
    print(f"Wrote {out_pdf}")


if __name__ == "__main__":
    build()
    to_pdf()
