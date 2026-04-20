"""Build the v2.0 DOCX of the cross-domain-techniques paper.

Produces ``docs/papers/cross-domain-techniques-v2.docx`` — the academic
revision of the Zenodo v1.0 PDF, now including:
- Real Evaluation section with numbers sourced directly from
  ``v041_public_datasets.json``.
- SHIPPED vs PROPOSED markers per technique.
- Limitations section that owns the tradeoffs.
- Vetted citations (research-agent verified April 2026).

The prose is written here, not copied from the markdown source. Tables
are native Word tables (not ASCII). Reproduce any time by running this
script; do NOT hand-edit the generated DOCX.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).parent.parent.parent
DATA_JSON = REPO_ROOT / "docs" / "papers" / "evaluation" / "v041_public_datasets.json"
OUT_DOCX = REPO_ROOT / "docs" / "papers" / "cross-domain-techniques-v2.docx"


# ---------------------------------------------------------------------------
# Small helpers
# ---------------------------------------------------------------------------


def _set_style_font(doc: Document, style_name: str, *, size: int | None = None, bold: bool | None = None) -> None:
    style = doc.styles[style_name]
    style.font.name = "Calibri"
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold


def _apply_defaults(doc: Document) -> None:
    """Set sensible font defaults across heading / body styles."""
    # Body
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Headings
    _set_style_font(doc, "Heading 1", size=18, bold=True)
    _set_style_font(doc, "Heading 2", size=14, bold=True)
    _set_style_font(doc, "Heading 3", size=12, bold=True)


def _add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: int | None = None, align: str | None = None) -> None:
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)


def _set_cell_bg(cell, color_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        _set_cell_bg(cell, "D9E1F2")
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _mono(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.8)


# ---------------------------------------------------------------------------
# Content assembly
# ---------------------------------------------------------------------------


def _format_pct(x: float) -> str:
    return f"{x:.3f}"


def build_document() -> Document:
    data = json.loads(DATA_JSON.read_text(encoding="utf-8"))
    datasets = data["datasets"]

    doc = Document()
    _apply_defaults(doc)

    # ----------------------- TITLE BLOCK -----------------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Beyond Pattern Matching")
    r.bold = True
    r.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Seven Cross-Domain Techniques for Prompt Injection Detection")
    r.bold = True
    r.font.size = Pt(16)

    _add_para(doc, "", size=4)  # spacer

    _add_para(doc, "Thamilvendhan Munirathinam", align="center", bold=True, size=12)
    _add_para(doc, "Independent Researcher  ·  prompt-shield project", align="center", italic=True, size=10)
    _add_para(doc, "Version 2.0 draft  ·  Revision date: 20 April 2026", align="center", size=10)
    _add_para(doc, "DOI (v1.0 anchor): 10.5281/zenodo.19644135", align="center", size=10)
    _add_para(doc, "Repository: github.com/mthamil107/prompt-shield", align="center", size=10)

    _add_para(doc, "", size=8)

    # ----------------------- ABSTRACT --------------------------
    doc.add_heading("Abstract", level=1)
    _add_para(
        doc,
        "Current open-source prompt-injection detectors converge on two architectural choices: "
        "regular-expression pattern matching and fine-tuned transformer classifiers. Both share failure "
        "modes that recent work has made concrete. Regular expressions miss paraphrased attacks. "
        "Fine-tuned classifiers are vulnerable to adaptive adversaries: a 2025 NAACL Findings study "
        "reported that eight published indirect-injection defenses were bypassed with greater than fifty "
        "percent attack-success rates under adaptive attacks, and a subsequent arXiv preprint extended "
        "that result to jailbreak defenses more broadly. Progress along the current axis of work appears "
        "to face diminishing returns against knowledgeable adversaries.",
    )
    _add_para(
        doc,
        "This work takes a different path. We propose seven detection techniques that each port a "
        "specific mechanism from a discipline outside large-language-model security: forensic "
        "linguistics, materials-science fatigue analysis, deception technology from network security, "
        "local-sequence alignment from bioinformatics, mechanism design from economics, spectral "
        "signal analysis with change-point detection from epidemiology, and taint tracking from "
        "compiler theory. Each mechanism produces a detection signal that is architecturally "
        "independent of both regular-expression matching and transformer classification, so the seven "
        "techniques compose with existing defenses rather than replacing them.",
    )
    _add_para(
        doc,
        "The v2.0 revision of this paper moves the contribution from proposal to partial empirical "
        "validation. Three of the seven techniques are implemented in the prompt-shield v0.4.1 "
        "release (Apache 2.0) and evaluated in a four-configuration ablation across six datasets: the "
        "public benchmarks deepset/prompt-injections, NotInject, LLMail-Inject, AgentHarm, and "
        "AgentDojo, and a synthetic indirect-injection benchmark released alongside this paper. The "
        "local-alignment detector lifts F1 on deepset from 0.033 to 0.378 (a thirty-four-point-six "
        "percentage-point improvement) with zero additional false positives. The stylometric "
        "detector adds eleven-point-one percentage points of F1 on the indirect-injection set while "
        "remaining silent on short inputs by design. The fatigue tracker is validated by a "
        "probing-campaign integration test: ten priming scans below the detection threshold cause a "
        "subsequent lower-confidence scan from the same source to be blocked. The remaining four "
        "techniques are described as design specifications and are left to future work. All claims, "
        "limitations, and reproduction scripts are released together with the paper.",
    )

    # ----------------------- 1. INTRODUCTION -------------------
    doc.add_heading("1. Introduction", level=1)

    _add_para(
        doc,
        "Prompt injection is the dominant offensive technique against large-language-model-integrated "
        "applications. A malicious input, whether supplied directly by an end user or smuggled through "
        "a retrieval-augmented-generation chunk, an email body, or a tool-call response, can steer the "
        "assistant model away from its operator-intended behavior. The attack class is mature enough "
        "to support a formalized taxonomy, public benchmarks, and a competitive defense research "
        "community, yet the defense literature remains narrow in its algorithmic palette.",
    )
    _add_para(
        doc,
        "A literature review of recent open-source detectors reveals an architectural convergence. "
        "Every widely-used defense tool we examined uses some combination of three mechanisms: a "
        "hand-written regular-expression pack for known attack patterns; a transformer classifier "
        "fine-tuned on labeled injection corpora; and a vector-similarity lookup against a store of "
        "historical malicious prompts. Each mechanism has well-understood weaknesses. Regular "
        "expressions miss semantic paraphrases and non-English attacks. Fine-tuned classifiers show "
        "strong in-distribution accuracy but, as recent adaptive-attack studies demonstrate, can be "
        "bypassed by knowledge-informed adversaries. Vector similarity cannot catch attacks that are "
        "semantically novel relative to the stored corpus.",
    )
    _add_para(
        doc,
        "The question this paper asks is whether meaningfully different detection signals can be "
        "recovered by porting mechanisms from disciplines that have solved structurally analogous "
        "problems in contexts unrelated to large-language-model security. We identify seven such "
        "mechanisms and present them as a coordinated palette of defenses, grouped by the detection "
        "signal they produce rather than by the attack class they target. Three techniques are "
        "implemented in the open-source prompt-shield library and empirically evaluated. The "
        "remaining four are described as design specifications pending implementation.",
    )

    doc.add_heading("1.1 Contributions", level=2)

    _add_para(
        doc,
        "The primary contributions of this v2.0 revision are:",
    )

    contrib = [
        (
            "Seven cross-domain detection mechanisms.",
            "We describe seven techniques, each rooted in a distinct scientific discipline "
            "with a self-contained mathematical or architectural basis: stylometric "
            "discontinuity, adversarial fatigue tracking, honeypot tool definitions, "
            "local-sequence alignment with a semantic substitution matrix, prediction-market "
            "ensemble scoring, perplexity spectral analysis, and runtime taint tracking for "
            "agent pipelines. Each is motivated by a specific failure mode of the existing "
            "regex-plus-classifier paradigm.",
        ),
        (
            "Three implementations with empirical validation.",
            "The local-alignment detector (d028), the adversarial fatigue tracker, and the "
            "stylometric discontinuity detector (d027) are released in prompt-shield v0.4.1 "
            "under the Apache 2.0 license. Each is accompanied by unit and integration tests "
            "(one hundred and three tests in total across the three techniques) and a public "
            "reproduction harness that regenerates every number in Section 5 from a single "
            "command invocation.",
        ),
        (
            "A four-configuration ablation across six datasets.",
            "We benchmark the three shipped techniques on five public datasets "
            "(deepset/prompt-injections, NotInject, LLMail-Inject Phase 1, AgentHarm, and "
            "AgentDojo v1.2.1) and on a new synthetic indirect-injection benchmark "
            "generated by a reproducible template-based builder also released alongside the "
            "paper. Every ablation cell reports precision, recall, F1, and false-positive "
            "rate directly from the benchmark JSON.",
        ),
        (
            "Honest limitations section.",
            "Section 5.4 enumerates the tradeoffs that a peer-review audience is likely to "
            "challenge: the NotInject false-positive-rate regression under d028, the "
            "saturation of LLMail-Inject under the pre-existing regex pack, the orthogonal "
            "behavior of AgentHarm to all three shipped techniques, and the self-evaluation "
            "risk of the synthetic indirect-injection benchmark. Each is flagged with a "
            "concrete next-revision action item.",
        ),
    ]
    for head, body in contrib:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(head + " ")
        r.bold = True
        p.add_run(body)

    _add_para(
        doc,
        "The remainder of the paper is structured as follows. Section 2 reviews the current "
        "prompt-injection defense landscape and its known failure modes. Section 3 defines the threat "
        "model. Section 4 presents the seven cross-domain techniques. Section 5 reports the empirical "
        "evaluation of the three shipped detectors. Section 6 concludes with a sequencing for the "
        "remaining four techniques and a roadmap for the next revision.",
    )

    # ----------------------- 2. BACKGROUND ---------------------
    doc.add_heading("2. Background and Related Work", level=1)

    doc.add_heading("2.1 The convergence problem", level=2)
    _add_para(
        doc,
        "Formal treatment of prompt injection as a measurable attack class dates to Liu et al. "
        "(USENIX Security 2024), who introduced a benchmark of attack-and-defense pairs and a "
        "taxonomy that has since been adopted by most subsequent work. In the two years since, "
        "the defense literature has consolidated around three patterns.",
    )
    _add_para(
        doc,
        "The first pattern is rule-based detection. Early open-source tools such as LLM Guard and "
        "prompt-shield itself in its v0.3 releases couple curated regular-expression packs with "
        "severity labels and aggregation logic. Commercial products in this family, including "
        "Lakera's PINT (April 2024) and the Rebuff project (archived May 2025), share the same "
        "architecture. The strength of rule-based detection is high precision on known attack "
        "vocabularies and fast execution. The weakness is that paraphrased or synonym-substituted "
        "attacks bypass the patterns verbatim.",
    )
    _add_para(
        doc,
        "The second pattern is classifier-based detection. Meta's Llama Prompt Guard 2 (2025, 22M "
        "and 86M parameter variants; no peer-reviewed paper) and Deepset's DeBERTa-v3 injection "
        "classifier are representative. ProtectAI's PIGuard, published at ACL 2025, explicitly "
        "addresses the over-defense problem by releasing the NotInject benchmark alongside a "
        "classifier calibrated for benign queries that contain attack-adjacent vocabulary. "
        "Classifiers absorb paraphrased attacks better than rules but introduce substantial "
        "false-positive rates on legitimate queries about the attack class itself, and are the "
        "specific target of the adaptive-attack research summarized in the next subsection.",
    )
    _add_para(
        doc,
        "The third pattern is meta-defense: model-integrated signaling. Spotlighting (Hines et al., "
        "arXiv:2403.14720) encodes context data with delimiters or simple transformations so the "
        "model can distinguish trusted instructions from untrusted context. SelfDefend (Wu et al., "
        "arXiv:2406.05498) deploys a shadow large-language-model whose sole job is to classify the "
        "primary model's inputs. MELON (Zhu et al., ICML 2025, arXiv:2502.05174) executes the "
        "primary model twice and blocks when the two runs diverge on tool invocations. These "
        "meta-defenses introduce latency but can catch attacks the primary detector misses.",
    )

    doc.add_heading("2.2 Adaptive attacks as the current frontier", level=2)
    _add_para(
        doc,
        "Zhan et al. (NAACL 2025 Findings, arXiv:2503.00061) evaluated eight published "
        "indirect-injection defenses under adaptive attacks and reported that all eight were bypassed "
        "with attack-success rates exceeding fifty percent. Nasr et al. (arXiv:2510.09023, "
        "OpenReview 7B9mTg7z25, submission status pending) extended the result to jailbreak "
        "defenses and reported greater than ninety percent attack-success rates against twelve "
        "published defenses. Both results target the same underlying weakness: when the adversary "
        "has model-gradient or decision-boundary access, any single-mechanism defense can be "
        "systematically evaded.",
    )
    _add_para(
        doc,
        "The methodological implication is that single-axis defense research is subject to "
        "diminishing returns. A detector that is ten percent stronger on a fixed benchmark is "
        "typically not measurably harder to adaptively evade than its predecessor. The defenses "
        "that remain effective under adaptive attacks tend to exploit properties the adversary "
        "cannot cheaply manipulate, such as context provenance or cross-model disagreement. This "
        "paper argues that introducing mechanistically independent detection signals creates "
        "additional invariants that an adversary must simultaneously evade, which is strictly "
        "harder than evading a stronger version of a single signal.",
    )

    doc.add_heading("2.3 Agentic-attack benchmarks and tool integrity", level=2)
    _add_para(
        doc,
        "The agent-integrated deployment pattern introduces attack surface beyond the text-input "
        "boundary. AgentDojo (Debenedetti et al., NeurIPS 2024, arXiv:2406.13352) formalizes "
        "indirect injection against tool-using agents with a suite of benign user tasks paired "
        "with injection tasks embedded in tool responses. Agent Security Bench (Zhang et al., "
        "ICLR 2025, arXiv:2410.02644) scales the formalism to ten scenarios with eleven defenses "
        "and four hundred tools. AgentHarm (Andriushchenko et al., ICLR 2025, arXiv:2410.09024) "
        "measures harmfulness of agent task completion rather than injection success directly; as "
        "we show in Section 5, this distinction matters for detector evaluation.",
    )
    _add_para(
        doc,
        "Two recent works target the same agentic-attack surface from opposite directions. "
        "Information flow control for agents is developed in FIDES (Costa et al., Microsoft "
        "Research, arXiv:2505.23643), which attaches confidentiality and integrity labels to "
        "model and tool inputs. Static taint analysis for large-language-model-integrated "
        "applications is developed in TaintP2X (He et al., ICSE 2026), which treats the "
        "assistant's prompt assembly as a dataflow graph and checks reachability from "
        "untrusted sources to sensitive sinks. Together these works establish that "
        "architectural-integrity techniques from compiler security are viable in the "
        "large-language-model context, a premise our Section 4.7 proposal builds on.",
    )

    # ----------------------- 3. THREAT MODEL --------------------
    doc.add_heading("3. Threat Model", level=1)
    _add_para(
        doc,
        "We consider a deployed assistant pipeline consisting of a system prompt, a user input "
        "channel, zero or more retrieval-augmented context chunks, and optionally a tool-calling "
        "capability that returns structured results back into the model context. An adversary "
        "controls at least one of the non-system channels: the direct user input, an RAG source "
        "they can populate, an email or document body the assistant will summarize, or a tool "
        "response the attacker can craft by compromising an upstream data source.",
    )
    _add_para(
        doc,
        "The adversary's goal is to induce the model to ignore its operator-intended behavior, "
        "either by exfiltrating the system prompt, by invoking tools the operator did not "
        "authorize, or by emitting content the operator explicitly forbade. The adversary may "
        "paraphrase attacks, substitute synonyms, insert filler words, obfuscate with encoding "
        "or unicode homoglyphs, pose attacks as hypothetical questions or academic research, "
        "split payloads across multiple turns, or exploit tool-response channels to smuggle "
        "attacks into otherwise-trusted contexts.",
    )
    _add_para(
        doc,
        "We treat the adversary as adaptive in the sense of Zhan et al. and Nasr et al.: they "
        "have access to the defense source code, the detector thresholds, and the published "
        "benchmark results. They do not have access to per-deployment state such as the "
        "fatigue tracker's per-source EWMA (Section 4.2), the operator's canary tokens, or the "
        "self-learning vault's accumulated history. Techniques that depend on such "
        "per-deployment state are strictly stronger than techniques that depend only on the "
        "released code.",
    )
    _add_para(
        doc,
        "We do not model attacks that subvert the operator or the model-weight pipeline itself. "
        "Supply-chain risks, training-data poisoning, and fine-tune-time backdoors are out of "
        "scope. Also out of scope are side-channel attacks against the underlying inference "
        "infrastructure (timing, memory disclosure, GPU isolation). Our threat model concerns "
        "adversaries operating through the normal user-facing input channels of an already-"
        "provisioned pipeline.",
    )

    return doc, datasets


def build_and_continue(doc: Document, datasets: dict) -> None:
    """Sections 4-References appended to the document built above."""

    # ----------------------- 4. TECHNIQUES ---------------------
    doc.add_heading("4. Seven Cross-Domain Techniques", level=1)
    _add_para(
        doc,
        "Each subsection below identifies the source discipline, describes the detection mechanism "
        "in enough detail for replication, and reports the current implementation status. The three "
        "shipped techniques (Sections 4.1, 4.2, 4.4) include references to the source code and unit "
        "tests in prompt-shield v0.4.1; their empirical evaluation is in Section 5. The four remaining "
        "techniques (Sections 4.3, 4.5, 4.6, 4.7) are presented as design specifications.",
    )

    # --- 4.1 Stylometric ---
    doc.add_heading("4.1 Stylometric Discontinuity Detection", level=2)
    _add_para(doc, "Status: SHIPPED in prompt-shield v0.4.1 (Apache 2.0). Source: src/prompt_shield/detectors/d027_stylometric_discontinuity.py.", italic=True)
    _add_para(
        doc,
        "Forensic linguistics has studied authorship attribution since the Federalist Papers "
        "analysis of Mosteller and Wallace in the 1960s. The stylometric tradition holds that "
        "individual authors leave measurable signatures in their writing — distributions of "
        "function-word frequency, sentence-length variance, vocabulary richness — that survive "
        "paraphrase, topic variation, and deliberate obfuscation. Applied to the task of "
        "detecting multi-author documents, the contemporary Plagiarism Analysis shared task PAN "
        "(Bevendorff et al., CLEF 2024) formalizes a sliding-window approach: partition a document "
        "into overlapping token windows, compute a stylometric feature vector per window, and "
        "measure divergence between adjacent windows to locate authorship boundaries.",
    )
    _add_para(
        doc,
        "An indirect prompt injection necessarily introduces a second author into a document. "
        "The legitimate author, typically a human writing an email or a content-management "
        "system emitting an HTML page, produces text in one stylistic distribution. The attacker "
        "producing the injected payload writes in a different distribution, characterized by "
        "imperative mood, high uppercase-character frequency, and directive vocabulary. The "
        "stylometric approach recasts indirect-injection detection as a well-studied "
        "style-change-boundary problem for which there is a decades-long feature engineering "
        "literature to draw from.",
    )
    _add_para(
        doc,
        "The detector extracts nine features per fifty-token window with twenty-five-token "
        "stride: function-word frequency (closed-class words such as the, is, of); average word "
        "length; average sentence length; stylistic punctuation density; hapax legomena ratio "
        "(singletons divided by vocabulary size); Yule's K statistic for vocabulary richness; "
        "imperative-verb ratio over a closed list of attack-adjacent imperatives; uppercase "
        "character ratio; and all-capital-word ratio for tokens of length three or more. Each "
        "feature is scaled to a [0, 1] range, the resulting nine-dimensional vector is "
        "normalized to a probability mass, and the Jensen-Shannon divergence "
        "(Lin, IEEE TIT 1991) is computed between every adjacent window pair. The maximum "
        "divergence across the input is compared against a calibrated threshold.",
    )
    _add_para(
        doc,
        "The detector is silent on inputs shorter than one hundred tokens because the "
        "per-window feature estimates are unstable on smaller samples; this is a deliberate "
        "engineering choice, not a limitation that affects the detection claim. On our "
        "synthetic indirect-injection benchmark (Section 5) the detector lifts the F1 score "
        "of a twenty-six-detector regex baseline from 0.889 to 1.000 with zero false "
        "positives on benign long-form distractors.",
    )
    _add_para(
        doc,
        "Prior application of stylometric features to large-language-model security has, to "
        "our knowledge, been limited to the authorship-attribution direction: distinguishing "
        "human-written from machine-generated text (Opara 2024, arXiv:2405.10129; "
        "arXiv:2507.00838). Reversing the task to detect adversarial-author injection in "
        "otherwise-benign documents appears to be new.",
    )

    # --- 4.2 Fatigue ---
    doc.add_heading("4.2 Adversarial Fatigue Tracking", level=2)
    _add_para(doc, "Status: SHIPPED in prompt-shield v0.4.1 (opt-in, Apache 2.0). Source: src/prompt_shield/fatigue/tracker.py.", italic=True)
    _add_para(
        doc,
        "Structural engineering has long distinguished between single-cycle failure and fatigue "
        "failure. A beam that does not break under a one-time load of a given magnitude may still "
        "fail under repeated loading at a lower magnitude, because microscopic cracks accumulate "
        "below the yield threshold and eventually propagate into macroscopic fractures. Basquin's "
        "1910 stress-life curves and the broader fatigue-analysis literature (Suresh, Cambridge "
        "University Press, second edition 1998) model this accumulation as a function of both "
        "the stress amplitude and the number of cycles.",
    )
    _add_para(
        doc,
        "A sophisticated adversary probing a prompt-injection detector exhibits the same "
        "statistical signature: a sustained sequence of inputs that each individually score "
        "just below the detection threshold, iteratively searching for the minimum-effort "
        "bypass. A stateless detector that evaluates each input in isolation cannot "
        "distinguish this probing campaign from a burst of benign near-boundary queries. A "
        "stateful detector that models the rate of near-threshold scores per source can.",
    )
    _add_para(
        doc,
        "We define a near-miss for a given (source, detector) pair as a scan whose confidence "
        "lies in the interval [threshold minus delta, threshold), where delta is a "
        "configuration parameter. The tracker maintains an exponentially-weighted moving "
        "average (Roberts 1959; see Tsai et al., IET Information Security 2024, for "
        "contemporary intrusion-detection use) of the near-miss indicator over successive "
        "scans. When the EWMA exceeds a configured trigger ratio and a minimum number of "
        "samples have been observed, the detector enters a hardened state for that source, "
        "in which the effective threshold is lowered by a configurable amount. The hardened "
        "state automatically releases after a cooldown interval with no further near-misses.",
    )
    _add_para(
        doc,
        "Structural epidemiology uses the same mathematical tools for outbreak detection. "
        "EWMA and CUSUM statistics are standard in syndromic-surveillance systems "
        "(Centers for Disease Control and Prevention; Vial et al., Emerging Infectious "
        "Diseases 26(9), 2020), where the task is to identify a shift in the baseline rate "
        "of a signal that is noisy on any single day. The cross-domain analogy that "
        "motivates the detector is therefore twofold: materials-science fatigue supplies "
        "the failure-model framing, and epidemiological surveillance supplies the "
        "detection-statistic choice.",
    )
    _add_para(
        doc,
        "The fatigue tracker is deliberately orthogonal to static single-sample benchmarks "
        "because its signal is temporal. We validate it end-to-end in Section 5.3 via a "
        "probing-campaign integration test: after ten priming scans from a single source at "
        "confidence 0.65 (base threshold 0.7), the eleventh scan from the same source at "
        "confidence 0.63 — strictly below every priming score and strictly below the "
        "un-hardened threshold — is blocked. A concurrent scan from a different source at "
        "the same confidence is unaffected, confirming per-source isolation. The tracker is "
        "opt-in and imposes zero overhead when disabled.",
    )

    # --- 4.3 Honeypot ---
    doc.add_heading("4.3 Honeypot Tool Definitions", level=2)
    _add_para(doc, "Status: PROPOSED. Implementation planned as the prompt_shield.honeypot module.", italic=True)
    _add_para(
        doc,
        "Network security has decades of experience with deception: a honeypot is a resource "
        "with no legitimate use case such that any interaction is definitionally adversarial. "
        "Recent work has extended the pattern to the large-language-model agent setting. "
        "Reworr and Volkov (Palisade Research, arXiv:2410.13919) deployed honeypots on the "
        "public internet to observe large-language-model-driven attackers and characterize "
        "their behavior. Pasquini et al. (arXiv:2410.20911) designed a defensive honeypot "
        "system that uses prompt injection against the attacker, demonstrating that the same "
        "attack surface can be inverted when the defender chooses the environment.",
    )
    _add_para(
        doc,
        "We propose extending deception to the tool-call layer of large-language-model "
        "agents. Agent frameworks typically expose a list of tool definitions to the "
        "model; we add canary tool definitions with names such as `get_admin_access`, "
        "`dump_database`, `reveal_system_prompt`, and `disable_safety_filters`. These tools "
        "have no legitimate invocation path. An agent pipeline that observes any call to a "
        "canary tool can treat the call as a confirmed indirect-injection event with one "
        "hundred percent precision: no benign query path leads to these tool names. The "
        "design is complementary to canary tokens embedded in system prompts (an older "
        "technique popularized by the Rebuff project, 2023) and to integrity-probe checks "
        "that periodically test whether the system prompt is still intact.",
    )
    _add_para(
        doc,
        "Validation requires an agent-harness simulator that exercises the tool-call "
        "boundary under both benign and adversarial workloads. Section 6 sequences this as "
        "the first remaining-technique implementation because the engineering cost is "
        "modest relative to the other three pending techniques.",
    )

    # --- 4.4 Smith-Waterman ---
    doc.add_heading("4.4 Local Sequence Alignment with a Semantic Substitution Matrix", level=2)
    _add_para(doc, "Status: SHIPPED in prompt-shield v0.4.1 as d028_sequence_alignment (Apache 2.0). Source: src/prompt_shield/detectors/d028_sequence_alignment.py.", italic=True)
    _add_para(
        doc,
        "Bioinformatics has the strongest continuous tradition of robust substring-matching "
        "research of any computing discipline. The Smith-Waterman algorithm (Smith and "
        "Waterman, J. Mol. Biol. 147:195-197, 1981) computes the optimal local alignment "
        "between two biological sequences in quadratic time using dynamic programming; the "
        "BLOSUM family of scoring matrices (Henikoff and Henikoff, PNAS 89:10915-10919, 1992) "
        "encodes the evolutionary cost of substituting one amino acid for another, allowing "
        "the algorithm to recognize homology in the presence of mutation.",
    )
    _add_para(
        doc,
        "A prompt-injection attack and its paraphrase share structural properties that map "
        "cleanly onto this framework. The attack has a characteristic word order "
        "(`ignore all previous instructions`); a paraphrase substitutes tokens within "
        "semantic equivalence classes (`disregard every prior directive`); a padded attack "
        "inserts filler words between structural anchors (`please just kindly ignore, if "
        "you would, all of the previous instructions`). Each of these transformations is "
        "analogous to a mutation, insertion, or deletion in a biological sequence, and "
        "Smith-Waterman with a semantically-weighted substitution matrix recovers the "
        "alignment despite the transformation.",
    )
    _add_para(
        doc,
        "The d028 detector maintains a curated database of approximately one hundred and "
        "eighty attack sequences across thirteen attack categories, drawn from the prompt-"
        "shield regex patterns of detectors d001 through d020. The substitution matrix has "
        "fifteen synonym groups covering the attack-adjacent vocabulary (ignore-family, "
        "instruction-family, reveal-family, system-family, etc.). The scoring function "
        "rewards exact matches with +3, same-group synonym matches with +2, "
        "unrelated-token mismatches with -1, and gap insertions with -1. A strict-greater-"
        "than comparison against a normalized threshold of 0.6 prevents false positives "
        "from generic English prefixes such as \"show me the\" that happen to align with "
        "the opening of an attack needle.",
    )
    _add_para(
        doc,
        "On the deepset/prompt-injections benchmark — the canonical academic dataset for "
        "prompt-injection detection — the twenty-six-detector regex baseline catches one "
        "attack out of sixty (F1 0.033) because the samples are dominated by paraphrased "
        "attacks that the regex patterns miss verbatim. The d028 detector catches fourteen "
        "of sixty (F1 0.378), a thirty-four-point-six percentage-point F1 improvement, "
        "with zero additional false positives.",
    )
    _add_para(
        doc,
        "To our knowledge, no prior work applies semantically-weighted local sequence "
        "alignment to prompt-injection detection. The CRAN text.alignment package applies "
        "Smith-Waterman to plain-text document comparison but uses character-level "
        "matching without a semantic matrix.",
    )

    # --- 4.5 Prediction Market ---
    doc.add_heading("4.5 Prediction-Market Ensemble Scoring", level=2)
    _add_para(doc, "Status: PROPOSED. Implementation blocked on data-model work; see Section 6.", italic=True)
    _add_para(
        doc,
        "Mechanism design from economics offers mature solutions to the information-aggregation "
        "problem that ensemble scoring solves in an ad-hoc way. A prediction market, equipped "
        "with a proper scoring rule such as Hanson's Logarithmic Market Scoring Rule "
        "(Journal of Prediction Markets, 2007), aggregates beliefs from heterogeneous "
        "participants into a single calibrated probability, weighting each participant by their "
        "historical accuracy (measured, for example, by Brier score).",
    )
    _add_para(
        doc,
        "Current ensemble scoring in prompt-shield and comparable tools uses a maximum-"
        "confidence-plus-bonus rule that ignores detector reliability and double-counts "
        "correlated signals. Replacing this rule with an LMSR market where each detector's "
        "bet is weighted by its historical Brier score over labeled scan history would produce "
        "better-calibrated probabilities and automatically down-weight unreliable detectors. "
        "Implementation requires a new persistence schema for per-detector confidence history "
        "and a ground-truth labeling path. The migration carries production risk because it "
        "touches the core scoring path; we plan a mandatory shadow-mode validation gate in "
        "which both legacy and market scoring run side-by-side with the legacy score driving "
        "actions until multi-week agreement is established.",
    )

    # --- 4.6 Spectral ---
    doc.add_heading("4.6 Perplexity Spectral Analysis", level=2)
    _add_para(doc, "Status: PROPOSED. Implementation requires an optional transformers dependency for GPT-2-small (124M params).", italic=True)
    _add_para(
        doc,
        "Signal processing and epidemiology share a concern with change-point detection in "
        "noisy sequences. A benign document produces a smooth, low-frequency perplexity "
        "signal when scored token-by-token by a reference language model. An injected "
        "payload with different vocabulary, syntax, and intent creates a high-frequency "
        "spike. We propose computing the discrete Fourier transform of the per-token "
        "perplexity series and flagging inputs with abnormally high high-frequency energy "
        "ratio, complemented by cumulative-sum change-point detection "
        "(Page, Biometrika 1954; Vial et al., CDC Emerging Infectious Diseases 2020) to "
        "localize the boundary between benign prose and injected payload.",
    )
    _add_para(
        doc,
        "The technique targets embedded or sandwich-style attacks where the payload is "
        "surrounded by benign cover text, a class that rule-based and classifier-based "
        "defenses have difficulty with when the cover text dominates the input. The main "
        "implementation cost is the optional dependency on a small reference model; we "
        "plan to use GPT-2-small (124M parameters) with lazy loading, matching the "
        "pattern already used for the existing d022 semantic classifier.",
    )

    # --- 4.7 Taint Tracking ---
    doc.add_heading("4.7 Runtime Taint Tracking for Agent Pipelines", level=2)
    _add_para(doc, "Status: PROPOSED. Design sketch below; empirical validation requires an agent-pipeline fixture not yet constructed.", italic=True)
    _add_para(
        doc,
        "Compiler security has studied taint analysis since the 1970s: data values are "
        "labeled with provenance, label propagation follows dataflow, and a policy violation "
        "is raised when tainted data reaches a privileged sink without passing through a "
        "sanitizer. The pattern applies directly to agent pipelines: system-prompt data is "
        "trusted, retrieval context is semi-trusted, user input is untrusted, and tool "
        "responses are untrusted unless the tool is known-safe. A tool call invoked with "
        "arguments that trace back to an untrusted source without passing through an "
        "explicit sanitizer is a policy violation.",
    )
    _add_para(
        doc,
        "We propose a runtime variant of the static analysis developed in TaintP2X "
        "(He et al., ICSE 2026): a `TaintedString` wrapper that carries provenance metadata "
        "through string operations, a sink-validation hook on the tool-call boundary, and a "
        "`TaintViolation` exception that callers can catch. The runtime variant is strictly "
        "weaker than static analysis in the formal sense (it only catches violations on the "
        "actual execution path) but strictly more practical: it requires no whole-program "
        "analysis and operates on the live object graph. Information Flow Control for AI "
        "agents (Costa et al., FIDES, arXiv:2505.23643) establishes that this framing is "
        "viable in the large-language-model setting.",
    )

    # ----------------------- 5. EVALUATION ---------------------
    doc.add_heading("5. Evaluation", level=1)

    doc.add_heading("5.1 Methodology", level=2)
    _add_para(
        doc,
        "We evaluate the three shipped detectors via a four-configuration ablation across "
        "six datasets. The four configurations are: a baseline twenty-six-detector regex "
        "pack identical to the prompt-shield v0.3.3 release (with the d022 semantic "
        "classifier deliberately disabled so the ablation isolates the regex-plus-novel "
        "contribution); that baseline plus d028 only; the baseline plus d027 only; and the "
        "baseline plus both novel detectors. The fatigue tracker is evaluated separately "
        "in Section 5.3 because it signals on temporal sequences rather than on individual "
        "samples.",
    )
    _add_para(
        doc,
        "The detection rule on every sample is: the input is counted as detected when the "
        "engine returns an action in {block, flag} or when the overall risk score is "
        "greater than or equal to 0.5. This matches the rule used by the repository's "
        "existing benchmark harness and by the v0.3.3 baseline recorded in "
        "tests/baseline_v0.3.3.txt. The full scan threshold is 0.7 in every configuration.",
    )
    _add_para(
        doc,
        "We use six datasets. Five are public and standard in the prompt-injection "
        "literature: deepset/prompt-injections (116 samples, test split), NotInject "
        "(339 benign samples across all three splits), microsoft/llmail-inject-challenge "
        "(a one-thousand-sample subset of the Phase 1 submissions), AgentHarm (the test_public "
        "split of both the harmful and harmless_benign configurations, 352 samples total), "
        "and AgentDojo v1.2.1 (132 tasks total, extracted from the agentdojo Python package). "
        "The sixth is a new synthetic indirect-injection benchmark (80 samples), generated by "
        "the reproducible builder docs/papers/evaluation/build_indirect_injection_benchmark.py "
        "and released alongside the paper. The benchmark combines six benign genre templates "
        "(financial report, team email, news article, research paper, technical "
        "documentation, analyst memo) with five egregious payload styles (uppercase-directive, "
        "credential exfiltration, system-override, admin-password request, safety-disable "
        "directive) to produce thirty positive samples; fifty negative samples use the same "
        "templates with benign mid-paragraph bridges. The seed is fixed at 20260420 for "
        "deterministic regeneration.",
    )

    doc.add_heading("5.2 Results", level=2)
    _add_para(
        doc,
        "Table 1 reports the full four-configuration ablation across all six datasets. "
        "Every number is sourced directly from v041_public_datasets.json produced by a "
        "single deterministic run of run_public_datasets.py; reproduction is documented in "
        "Section 5.5.",
    )

    # Build the big results table
    headers = ["Dataset (n)", "Config", "TP", "TN", "FP", "FN", "Prec.", "Recall", "F1", "FPR"]
    rows: list[list[str]] = []
    ordered = [
        ("deepset/prompt-injections", "deepset"),
        ("leolee99/NotInject", "NotInject"),
        ("microsoft/llmail-inject-challenge (Phase1, 1000 subset)", "LLMail-Inject"),
        ("ai-safety-institute/AgentHarm", "AgentHarm"),
        ("ethz-spylab/agentdojo (v1.2.1)", "AgentDojo"),
        ("synthetic/indirect-injection", "Synthetic indirect-injection"),
    ]
    for key, short in ordered:
        ds = datasets[key]
        n = ds["sample_count"]
        for cfg_key in ("baseline", "plus_d028", "plus_d027", "plus_d027_d028"):
            m = ds["configs"][cfg_key]
            rows.append([
                f"{short} ({n})",
                cfg_key,
                str(m["tp"]),
                str(m["tn"]),
                str(m["fp"]),
                str(m["fn"]),
                _format_pct(m["precision"]),
                _format_pct(m["recall"]),
                _format_pct(m["f1"]),
                _format_pct(m["fpr"]),
            ])
    _add_table(doc, headers, rows)
    _add_para(doc, "Table 1. Four-configuration ablation across six datasets. Precision, recall, F1 and FPR computed at the standard detection rule (action ∈ {block, flag} or overall_risk_score ≥ 0.5).", italic=True, size=9)

    doc.add_heading("5.3 Per-technique interpretation", level=2)

    _add_para(
        doc,
        "Local sequence alignment (d028). On deepset the detector lifts recall from "
        "one-in-sixty to fourteen-in-sixty (a 14× improvement), and F1 from 0.033 to "
        "0.378, with zero additional false positives. This is the strongest single-technique "
        "result in the evaluation and is consistent with the design intent: deepset is "
        "dominated by paraphrased attacks and synonym substitutions that the regex pack "
        "cannot see verbatim but that the substitution matrix recognizes. On NotInject the "
        "detector introduces ten false positives (FPR 0.9% → 3.8%), which we flag and "
        "discuss in Section 5.4. On LLMail-Inject the detector adds two true positives "
        "(recall 0.978 → 0.980), a saturation effect: the benchmark is dominated by direct "
        "instruction-override attacks that the regex pack already catches. On AgentHarm the "
        "detector moves no metrics, a consequence of the benchmark measuring agent-task "
        "harmfulness rather than injection detection. On AgentDojo the detector adds one "
        "true positive and three false positives, a net neutral result in F1 terms. On the "
        "synthetic indirect-injection set the detector closes the remaining six-sample gap "
        "from 0.889 to 1.000.",
    )
    _add_para(
        doc,
        "Stylometric discontinuity (d027). By construction the detector is silent on inputs "
        "shorter than one hundred tokens; because the five public datasets are dominated by "
        "shorter samples the detector produces no movement on any of their metrics. On the "
        "synthetic indirect-injection set — where every sample is a ≥150-token document — "
        "the detector lifts F1 from 0.889 to 1.000 with zero false positives, matching the "
        "d028 result on the same set. The two detectors share the same six-sample residual "
        "on the indirect-injection set because both the stylometric break signal and the "
        "alignment signal fire on the same uppercase-directive payload structure. Their "
        "orthogonality manifests across datasets rather than within any one dataset: d028 "
        "contributes on deepset where d027 is silent, and d027 would contribute on a "
        "benchmark of long documents where d028's attack-sequence database is less useful.",
    )
    _add_para(
        doc,
        "Adversarial fatigue. Because the detector signals on temporal sequences rather "
        "than individual samples, and because the five public datasets present each sample "
        "as an independent scan, a static-dataset F1 delta is not a meaningful measure. The "
        "detector is validated end-to-end by the integration test "
        "tests/fatigue/test_engine_integration.py::test_hardening_catches_next_near_miss. "
        "In that test, ten scans from a single source at confidence 0.65 (base threshold "
        "0.7) pass individually, the per-(source, detector) EWMA of the near-miss indicator "
        "crosses the trigger ratio of 0.3, and the pair enters the hardened state. The "
        "eleventh scan from the same source at confidence 0.63 — strictly below every "
        "priming score and strictly below the un-hardened threshold — is then blocked, "
        "because the effective threshold has been lowered by the configured hardening "
        "offset. A concurrent scan from a different source at the same confidence 0.63 "
        "passes, confirming per-source isolation.",
    )

    doc.add_heading("5.4 Limitations and threats to validity", level=2)

    limits = [
        (
            "NotInject false-positive rate regression.",
            "d028 introduces ten false positives on the 339 benign samples of NotInject, "
            "moving the FPR from 0.9% (three false positives in the baseline) to 3.8% "
            "(thirteen). The regression is a consequence of the substitution matrix's "
            "permissive synonym grouping: benign queries that pair reveal-family verbs with "
            "instruction-family nouns (for example, \"show me the instructions for a "
            "chemistry demonstration\") accumulate enough alignment score to cross the "
            "detector threshold. Threshold tuning from 0.60 to 0.63 plus dampening of the "
            "show/reveal synonym group are planned and will appear in the next revision.",
        ),
        (
            "LLMail-Inject saturation.",
            "The baseline regex pack already recalls 97.8% of LLMail-Inject attacks "
            "because the benchmark is dominated by direct instruction-override phrasings "
            "that the pack targets explicitly. The 0.2 percentage-point F1 improvement "
            "under d028 is honest but small, and should not be cited as the headline "
            "result. LLMail-Inject is most useful as a robustness-to-variation check "
            "rather than as an improvement benchmark.",
        ),
        (
            "AgentHarm orthogonality.",
            "All four configurations achieve the same F1 of 0.319 and FPR of 31.8% on "
            "AgentHarm because the benchmark measures agent-task harmfulness rather than "
            "injection content. The harmful prompts invoke legitimate action-verb + "
            "admin-noun combinations that the baseline regex pack already flags, and "
            "adding injection-specific detectors does not change that behavior. The 31.8% "
            "FPR is itself an argument for the proposed taint-tracking technique "
            "(Section 4.7) which would distinguish benign administrative requests from "
            "adversarial tool invocations by provenance rather than by lexical content.",
        ),
        (
            "Synthetic benchmark self-evaluation risk.",
            "The indirect-injection benchmark was constructed with knowledge of what the "
            "d027 and d028 detectors measure. The 1.000 F1 numbers on that benchmark are "
            "therefore not held-out evaluation numbers in the strict sense. The "
            "benchmark's generator is deterministic and fully transparent — the template "
            "set, payload set, and random seed are all in the repository — but external "
            "validation on a held-out human-written indirect-injection benchmark is the "
            "honest next step. Until that exists, the 1.000 F1 numbers should be read as "
            "evidence that the detectors behave as designed on their target attack class, "
            "not as external-validation accuracy.",
        ),
        (
            "No adaptive-attack evaluation.",
            "The ICLR-2025 NAACL-Findings and contemporaneous arXiv:2510.09023 adaptive-"
            "attack methodologies are not yet applied to the shipped detectors. An adaptive "
            "adversary with knowledge of d028's substitution matrix could craft payloads "
            "whose synonyms deliberately fall outside the listed groups; an adversary "
            "aware of d027's feature set could suppress the uppercase and imperative-verb "
            "signals. We scope adaptive evaluation for v3.0 of this paper.",
        ),
        (
            "Agent Security Bench not included.",
            "We attempted to include ASB (Zhang et al., ICLR 2025) in the evaluation, but "
            "the benchmark requires its GitHub framework for execution and is not "
            "available on HuggingFace in a dataset form we could extract statically. A "
            "static extraction path analogous to what we did for AgentDojo is possible and "
            "is listed as future work.",
        ),
        (
            "Competitor comparison not rerun.",
            "The repository contains a comparison harness that runs ProtectAI DeBERTa v2, "
            "PIGuard, Meta Prompt Guard 2, Deepset DeBERTa v3 and prompt-shield on "
            "deepset/prompt-injections and NotInject. That harness has not been rerun "
            "with d028 and d027 enabled. The mechanical work is small and is listed as a "
            "v2.1 deliverable.",
        ),
    ]
    for head, body in limits:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(head + " ")
        r.bold = True
        p.add_run(body)

    doc.add_heading("5.5 Reproducibility", level=2)
    _add_para(doc, "Every number in Section 5.2 is produced by the following sequence, executable on Python 3.10 through 3.13:")
    _mono(doc, "git clone https://github.com/mthamil107/prompt-shield")
    _mono(doc, "cd prompt-shield")
    _mono(doc, "pip install -e \".[dev,all]\" agentdojo")
    _mono(doc, "python docs/papers/evaluation/build_indirect_injection_benchmark.py")
    _mono(doc, "python docs/papers/evaluation/run_public_datasets.py")
    _add_para(
        doc,
        "Outputs are written to v041_public_datasets.json (machine-readable) and "
        "v041_public_datasets.md (human-readable). Wall-clock runtime is approximately "
        "four minutes. The fatigue probing-campaign claim in Section 5.3 is reproduced "
        "by pytest tests/fatigue/ which runs in under one second. The full prompt-shield "
        "test suite comprises 868 tests on the current main branch and passes cleanly "
        "across Python 3.10, 3.11, 3.12, and 3.13 in continuous integration.",
    )

    # ----------------------- 6. CONCLUSIONS --------------------
    doc.add_heading("6. Conclusions and Future Work", level=1)
    _add_para(
        doc,
        "We presented seven cross-domain detection techniques for prompt injection and "
        "reported empirical results for three of them. The local-sequence-alignment "
        "detector delivers a thirty-four-point-six percentage-point F1 improvement on "
        "deepset/prompt-injections with zero additional false positives on that benchmark, "
        "while honestly introducing a 2.95 percentage-point false-positive-rate regression "
        "on NotInject that we have calibrated rather than hidden. The stylometric "
        "discontinuity detector lifts F1 by 11.1 percentage points on a synthetic "
        "indirect-injection benchmark released alongside the paper, while remaining silent "
        "on short inputs by design. The adversarial fatigue tracker is validated via a "
        "probing-campaign integration test that demonstrates the core behavioral claim: "
        "after a sustained burst of near-threshold scans from the same source, a "
        "subsequent lower-confidence scan from that source is blocked.",
    )
    _add_para(
        doc,
        "The four remaining techniques are sequenced by risk: honeypot tool definitions "
        "(Section 4.3) require a simulated-agent harness but pose no core-scoring risk; "
        "perplexity spectral analysis (Section 4.6) requires an optional ML dependency; "
        "runtime taint tracking (Section 4.7) requires an agent-pipeline fixture; "
        "prediction-market scoring (Section 4.5) touches the core scoring path and is "
        "scoped last, behind a mandatory shadow-mode gate.",
    )
    _add_para(
        doc,
        "The next revision of this paper will fold in: a head-to-head competitor "
        "comparison with Rebuff, Lakera PINT, Meta Prompt Guard 2, PIGuard, and Deepset "
        "DeBERTa v3 under the same four-configuration ablation; an adaptive-attack "
        "evaluation following Zhan et al. and Nasr et al.; a held-out indirect-injection "
        "benchmark composed of human-written documents with paraphrased payloads; and "
        "implementation and evaluation of at least one of the four remaining proposed "
        "techniques. Until those are in hand, the claims in this paper are scoped to the "
        "measurements reported in Section 5.",
    )

    # ----------------------- REFERENCES --------------------
    doc.add_heading("References", level=1)

    refs = [
        "Andriushchenko, M., Souly, A., et al. \"AgentHarm: A Benchmark for Measuring Harmfulness of LLM Agents.\" ICLR 2025. arXiv:2410.09024.",
        "Basquin, O. H. \"The Exponential Law of Endurance Tests.\" Proceedings of the American Society for Testing and Materials 10:625-630, 1910.",
        "Bevendorff, J., et al. \"Overview of PAN 2024: Multi-author Writing Style Analysis, Multilingual Text Detoxification, Oppositional Thinking Analysis, and Generative AI Authorship Verification.\" CLEF 2024. DOI 10.1007/978-3-031-71908-0_11.",
        "Costa, M., Köpf, B., et al. \"Securing AI Agents with Information-Flow Control.\" Microsoft Research, arXiv:2505.23643, 2025.",
        "Debenedetti, E., Zhang, J., Balunovic, M., Beurer-Kellner, L., Fischer, M., Tramèr, F. \"AgentDojo: A Dynamic Environment to Evaluate Prompt Injection Attacks and Defenses for LLM Agents.\" NeurIPS 2024 Datasets and Benchmarks. arXiv:2406.13352.",
        "He, X., Wang, B., Zhao, Y., Hou, X., Liu, J., Zou, H., Wang, H. \"TaintP2X: Detecting Taint-Style Prompt-to-Anything Injection Vulnerabilities in LLM-Integrated Applications.\" ICSE 2026 Research Track.",
        "Henikoff, S., Henikoff, J. G. \"Amino acid substitution matrices from protein blocks.\" Proceedings of the National Academy of Sciences 89(22):10915-10919, 1992. DOI 10.1073/pnas.89.22.10915.",
        "Hines, K., Lopez, G., Hall, M., Zarfati, F., Zunger, Y., Kiciman, E. \"Defending Against Indirect Prompt Injection Attacks With Spotlighting.\" arXiv:2403.14720, 2024.",
        "Hanson, R. \"Logarithmic Market Scoring Rules for Modular Combinatorial Information Aggregation.\" Journal of Prediction Markets 1(1):3-15, 2007.",
        "Li, H., Liu, Y., Zhang, C., Xiao, Y. \"PIGuard: Prompt Injection Guardrail via Mitigating Overdefense for Free.\" ACL 2025 Long Papers. aclanthology.org/2025.acl-long.1468.",
        "Lin, J. \"Divergence Measures Based on the Shannon Entropy.\" IEEE Transactions on Information Theory 37(1):145-151, 1991. DOI 10.1109/18.61115.",
        "Liu, Y., Jia, Y., Geng, R., Jia, J., Gong, N. Z. \"Formalizing and Benchmarking Prompt Injection Attacks and Defenses.\" USENIX Security 2024. arXiv:2310.12815.",
        "Nasr, M., Carlini, N., Sitawarin, C., et al. \"The Attacker Moves Second: Stronger Adaptive Attacks Bypass Defenses Against LLM Jailbreaks and Prompt Injections.\" arXiv:2510.09023, 2025. [Submission status pending; OpenReview 7B9mTg7z25.]",
        "Opara, C. \"StyloAI: Distinguishing AI-Generated Content with Stylometric Analysis.\" arXiv:2405.10129, 2024.",
        "Page, E. S. \"Continuous Inspection Schemes.\" Biometrika 41(1/2):100-115, 1954.",
        "Pasquini, D., Corti, E., Ateniese, G. \"Hacking Back the AI-Hacker: Prompt Injection as a Defense Against LLM-driven Cyberattacks.\" arXiv:2410.20911, 2024.",
        "Reworr, Volkov, D. \"LLM Agent Honeypot: Monitoring AI Hacking Agents in the Wild.\" Palisade Research, arXiv:2410.13919, 2024.",
        "Smith, T. F., Waterman, M. S. \"Identification of Common Molecular Subsequences.\" Journal of Molecular Biology 147:195-197, 1981. DOI 10.1016/0022-2836(81)90087-5.",
        "Suresh, S. Fatigue of Materials. Cambridge University Press, second edition, 1998. ISBN 9780521578479.",
        "Tsai, C.-W., et al. \"Using WPCA and EWMA Control Chart to Construct a Network Intrusion Detection Model.\" IET Information Security, 2024. DOI 10.1049/2024/3948341.",
        "Vial, F., et al. \"Assessing 3 Outbreak Detection Algorithms in an Electronic Syndromic Surveillance System.\" Emerging Infectious Diseases 26(9), US Centers for Disease Control and Prevention, 2020.",
        "Wu, X., Wang, R., et al. \"SelfDefend: LLMs Can Defend Themselves against Jailbreaking in a Practical Manner.\" arXiv:2406.05498, 2024.",
        "Zhan, Q., Fang, H., Panchal, A., Kang, D. \"Adaptive Attacks Break Defenses Against Indirect Prompt Injection Attacks on LLM Agents.\" NAACL 2025 Findings. arXiv:2503.00061.",
        "Zhang, J., Yu, R., et al. \"Agent Security Bench (ASB): Formalizing and Benchmarking Attacks and Defenses in LLM-based Agents.\" ICLR 2025. arXiv:2410.02644.",
        "Zhu, K., Yang, Y., Wang, R., Guo, Y., Wang, H. \"MELON: Provable Defense Against Indirect Prompt Injection Attacks in AI Agents.\" ICML 2025. arXiv:2502.05174.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style="List Number")
        p.paragraph_format.left_indent = Cm(0.6)

    # ----------------------- Artifact listing --------------------
    doc.add_heading("Appendix A. Released Artifacts", level=1)
    _add_para(
        doc,
        "All artifacts released alongside this paper are in the public repository at "
        "github.com/mthamil107/prompt-shield under the Apache 2.0 license. Relevant paths:",
    )
    artifacts = [
        ("src/prompt_shield/detectors/d028_sequence_alignment.py", "Smith-Waterman alignment detector (Section 4.4)."),
        ("src/prompt_shield/detectors/d027_stylometric_discontinuity.py", "Stylometric discontinuity detector (Section 4.1)."),
        ("src/prompt_shield/fatigue/tracker.py", "Adversarial fatigue tracker (Section 4.2)."),
        ("tests/detectors/test_d028_sequence_alignment.py", "35 tests covering alignment math, substitution matrix, detector behavior."),
        ("tests/detectors/test_d027_stylometric_discontinuity.py", "39 tests covering feature extraction, Jensen-Shannon math, detector behavior and fixtures."),
        ("tests/fatigue/test_tracker.py, test_engine_integration.py", "29 tests including the probing-campaign integration test cited in Section 5.3."),
        ("docs/papers/evaluation/run_public_datasets.py", "Reproduction harness for Table 1."),
        ("docs/papers/evaluation/build_indirect_injection_benchmark.py", "Generator for the synthetic indirect-injection benchmark."),
        ("docs/papers/evaluation/v041_public_datasets.json", "Machine-readable source of every number in Table 1."),
        ("docs/papers/evaluation/ANALYSIS.md", "Narrative analysis of per-dataset behavior with limitation framing."),
        ("tests/baseline_v0.3.3.txt", "Locked baseline for the regression gate used during the shipping of each detector."),
        ("CITATION.cff", "GitHub-native citation metadata."),
    ]
    for path, note in artifacts:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(path)
        r.font.name = "Consolas"
        r.font.size = Pt(10)
        p.add_run(f" — {note}")

    return doc


def main() -> None:
    doc, datasets = build_document()
    build_and_continue(doc, datasets)
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
